#!/usr/bin/env python3
"""Phase 3: compile the approved review draft into a Word .docx BD Profile.

Mechanical. No model call - pure parse of 02-assemble/output/profile-draft.md
into a styled Word document. Output is dated, so re-compiling preserves prior
versions.

House style: rather than rebuild the firm's chrome from scratch (which can never
match the custom M&Q banner, breadcrumb tag, logo header, footer and named
styles), this compiler opens the real M&Q template `03-compile/assets/base.docx`
as its base and writes the body into it. That guarantees byte-faithful branding:
  - the blue title banner ("Business Development Profile" + institution name,
    white on blue) and the "BD Profile" breadcrumb tag live in the template
    header; we only swap the institution-name placeholder and the footer date.
  - Heading 1/2 (bold blue), the Title/Subtitle banner styles, Calibri Light
    body, bullets (List Paragraph + the template's bullet numbering) all come
    from the template's own styles.xml, so nothing is approximated.

The body content rules (section order, leader subheads, headshots, the approval
gate) are unchanged from before. Headshots embed only when partner-approved in
.headshots.json; each is square-cropped and floated right with text wrapping,
matching the samples.

Usage:
  python 03-compile/build_docx.py
  python 03-compile/build_docx.py --in 02-assemble/output/profile-draft.md \
      --institution "Rhode Island School of Design" --date 2026-06-29

Writes 03-compile/output/BD_Profile_<institution>_<date>.docx
"""
import argparse
import datetime
import io
import json
import re
import sys
from copy import deepcopy
from pathlib import Path
from urllib import request, error

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "03-compile" / "assets"
BASE_DOC = ASSETS / "base.docx"
APPROVALS_FILE = ROOT / "02-assemble" / "output" / ".headshots.json"
sys.path.insert(0, str(ROOT))
from lib import profile  # noqa: E402

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install -r requirements.txt",
          file=sys.stderr)
    sys.exit(1)

GRAY = RGBColor(0x80, 0x80, 0x80)
HEADSHOT_IN = 2.08
PLACEHOLDER_NAME = "University of College"   # institution slot in the template
BULLET_NUM_GENERAL = 20                       # template bullet list (facts)
BULLET_NUM_LEADER = 42                         # template bullet list (experience)
SUBHEAD_RE = re.compile(r"^#{3}\s+(.+?)\s*$")
PHOTO_RE = re.compile(r"^Photo:\s*(.+?)\s*$", re.IGNORECASE)
CHART_RE = re.compile(r"^Chart:\s*(.+?)\s*$", re.IGNORECASE)
CHART_IN = 5.5                                 # rendered chart width in inches
MQ_BLUE = "#1F3A6E"
# Bare URLs and markdown [text](url). Bodies are plain Markdown with bare links,
# but accept the markdown form too. The bare pattern deliberately does not use
# \S+ alone: trailing sentence punctuation must stay OUT of the href.
URL_RE = re.compile(r"(?P<md>\[(?P<text>[^\]\n]+)\]\((?P<mdurl>https?://[^)\s]+)\))"
                    r"|(?P<bare>https?://[^\s<>\[\]{}|\\^\"]+)", re.IGNORECASE)
URL_TRAIL = ".,;:!?)]}'\"“”’"          # strip these off the end of a bare URL
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)  # Word's default hyperlink blue

_DISPLAY_BACK = {v: k for k, v in profile.DISPLAY.items()}


def _is_url(s):
    return s.lower().startswith(("http://", "https://"))


def canonical(heading):
    if heading in _DISPLAY_BACK:
        return _DISPLAY_BACK[heading]
    if heading in profile.ORIGINATION_HEADINGS:
        return profile.ORIGINATION_KEY
    return heading


def slug(name):
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_") or "Institution"


def title_from_draft(text, override):
    if override:
        return override
    m = re.search(r"^#\s+.*?BD Profile.*?:\s*(.+?)\s*$", text, re.MULTILINE)
    if m:
        return m.group(1).strip()
    m = re.search(r"^#\s+(.+?)\s*$", text, re.MULTILINE)
    return m.group(1).strip() if m else "Institution"


# ------------------------------------------------------ template chrome
def _set_para_text(p, text):
    """Collapse a paragraph's runs to a single run carrying `text` (keeps the
    first run's properties so the banner/breadcrumb styling survives)."""
    runs = p.findall(qn("w:r"))
    if not runs:
        return
    t = runs[0].find(qn("w:t"))
    if t is None:
        t = OxmlElement("w:t")
        runs[0].append(t)
    t.text = text
    t.set(qn("xml:space"), "preserve")
    for r in runs[1:]:
        p.remove(r)


def customize_chrome(doc, institution, stamp, breadcrumb):
    """Swap the institution-name placeholder in the header banner/breadcrumb,
    set the breadcrumb tag (BD Profile / Short BD Profile), and stamp the footer
    date. All edits are on the template's own header/footer parts."""
    def norm(s):
        return " ".join(s.split())
    for part in doc.part.package.iter_parts():
        pn = str(part.partname)
        if "header" in pn:
            el = part.element
            for p in el.iter(qn("w:p")):
                txt = norm("".join(t.text or "" for t in p.iter(qn("w:t"))))
                if txt == PLACEHOLDER_NAME:
                    _set_para_text(p, institution)
                elif txt == "BD Profile" and breadcrumb != "BD Profile":
                    _set_para_text(p, breadcrumb)
        elif "footer" in pn:
            for t in part.element.iter(qn("w:t")):
                if t.text and re.match(r"\d{4}-\d{2}-\d{2}", t.text.strip()):
                    t.text = stamp
    # Document-title metadata (Word's title property) also ships the placeholder.
    try:
        doc.core_properties.title = f"{breadcrumb}: {institution}"
        doc.core_properties.subject = institution
    except Exception:
        pass


def reset_body(doc):
    """Empty the template's placeholder body but keep the section that links to
    the header banner and footer; return that sectPr to append back at the end."""
    body = doc.element.body
    rich = None
    for sp in body.iter(qn("w:sectPr")):
        if sp.find(qn("w:headerReference")) is not None:
            rich = deepcopy(sp)
            break
    for child in list(body):
        body.remove(child)
    return body, rich


# ------------------------------------------------------------- body bits
def _small(paragraph, text, color=GRAY):
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = color
    return run


def _small_rich(doc, text):
    """A small gray provenance line that keeps markdown links clickable: text
    runs go 9pt gray, link runs go 9pt but keep their link styling."""
    p = doc.add_paragraph()
    _write_runs(p, text)
    for r in p.runs:                       # direct text runs
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    for r in p._p.iter(qn("w:r")):         # runs nested inside w:hyperlink
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
        if rPr.find(qn("w:sz")) is None:
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18")  # 9pt half-points
            rPr.append(sz)
    return p


def _hyperlink(paragraph, url, text):
    """Append a real, clickable hyperlink run. python-docx has no API for this:
    the URL must be registered as an external relationship and referenced from a
    w:hyperlink element, or Word renders it as dead text."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Style explicitly rather than trusting a "Hyperlink" style to exist in the
    # template; explicit run properties render the same everywhere.
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    rPr.append(color); rPr.append(underline)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _write_links(paragraph, text, bold=False):
    """Add `text` to `paragraph`, turning every URL into a clickable hyperlink.
    Plain runs carry `bold`; links keep their own styling."""
    def run(t):
        r = paragraph.add_run(t)
        if bold:
            r.bold = True
    pos = 0
    for m in URL_RE.finditer(text):
        if m.start() > pos:
            run(text[pos:m.start()])
        if m.group("md"):
            _hyperlink(paragraph, m.group("mdurl"), m.group("text"))
            pos = m.end()
            continue
        url = m.group("bare")
        trail = ""
        while url and url[-1] in URL_TRAIL:
            trail = url[-1] + trail
            url = url[:-1]
        if url:
            _hyperlink(paragraph, url, url)
        if trail:
            run(trail)   # sentence punctuation stays outside the link
        pos = m.end()
    if pos < len(text):
        run(text[pos:])
    return paragraph


def _write_bold(paragraph, text):
    """**spans** render bold; URLs become links inside either kind of span."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _write_links(paragraph, text[pos:m.start()])
        _write_links(paragraph, m.group(1), bold=True)
        pos = m.end()
    if pos < len(text):
        _write_links(paragraph, text[pos:])
    return paragraph


def _write_runs(paragraph, text):
    """Add `text` to `paragraph`. [verify]/[inferred] tags render highlighted so
    the review pass can't miss them; everything else gets bold + link handling."""
    from docx.enum.text import WD_COLOR_INDEX
    pos = 0
    for m in profile.TAG_RE.finditer(text):
        if m.start() > pos:
            _write_bold(paragraph, text[pos:m.start()])
        run = paragraph.add_run(m.group(0))
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        pos = m.end()
    if pos < len(text):
        _write_bold(paragraph, text[pos:])
    return paragraph


def _bullet(doc, text, level, num_id):
    p = doc.add_paragraph(style="List Paragraph")
    _write_runs(p, text)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(min(level, 2)))
    nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
    numPr.append(ilvl); numPr.append(nid)
    pPr.append(numPr)
    return p


def _square_crop(data):
    """Center-crop to a square, biased upward so faces are not beheaded.

    Sources are never square (institution headshots are typically 4:3 landscape
    or portrait), so a square is always a crop, never a resize.
    """
    from PIL import Image
    im = Image.open(io.BytesIO(data))
    if im.mode != "RGB":
        im = im.convert("RGB")
    w, h = im.size
    s = min(w, h)
    left = (w - s) // 2
    top = max(0, (h - s) // 4)   # bias up: heads sit above the vertical centre
    im = im.crop((left, top, left + s, top + s))
    out = io.BytesIO()
    im.save(out, format="JPEG", quality=88)
    out.seek(0)
    return out


def _floatify(run):
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline")) if drawing is not None else None
    if inline is None:
        return
    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    cNv = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))
    if extent is None or graphic is None:
        return
    anchor = OxmlElement("wp:anchor")
    for k, v in {"distT": "0", "distB": "91440", "distL": "114300",
                 "distR": "114300", "simplePos": "0",
                 "relativeHeight": "251658240", "behindDoc": "0", "locked": "0",
                 "layoutInCell": "1", "allowOverlap": "1"}.items():
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0")
    anchor.append(sp)
    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "column")
    al = OxmlElement("wp:align"); al.text = "right"; ph.append(al)
    anchor.append(ph)
    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "paragraph")
    off = OxmlElement("wp:posOffset"); off.text = "0"; pv.append(off)
    anchor.append(pv)
    anchor.append(extent)
    ee = OxmlElement("wp:effectExtent")
    for a in ("l", "t", "r", "b"):
        ee.set(a, "0")
    anchor.append(ee)
    wrap = OxmlElement("wp:wrapSquare"); wrap.set("wrapText", "bothSides")
    anchor.append(wrap)
    if docPr is not None:
        anchor.append(docPr)
    if cNv is not None:
        anchor.append(cNv)
    anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)


def _embed_headshot(doc, url, paragraph=None):
    """Float a headshot at `paragraph` (the leader's heading), not at a paragraph
    of its own. Anchoring to an existing paragraph keeps the image beside that
    leader's entry and adds no empty line to the flow."""
    try:
        req = request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        try:
            stream = _square_crop(data)
            size = {"width": Inches(HEADSHOT_IN), "height": Inches(HEADSHOT_IN)}
        except Exception:
            # Cropping failed (no Pillow, odd format). Embed as-is and let the
            # height scale: forcing a square onto a 4:3 source squashes the face.
            stream = io.BytesIO(data)
            size = {"width": Inches(HEADSHOT_IN)}
        host = paragraph if paragraph is not None else doc.add_paragraph()
        run = host.add_run()
        run.add_picture(stream, **size)
        _floatify(run)
        return True
    except (error.URLError, OSError, ValueError) as e:
        _small(doc.add_paragraph(), f"[headshot unavailable: {url} ({e.__class__.__name__})]")
        return False


_MULT = {"k": 1e3, "m": 1e6, "b": 1e9}


def _parse_money(s):
    """Parse a chart value like '$4.2M', '4,200,000', '$3.1 billion' -> float.
    Returns None if there is no usable number (e.g. a '$X' placeholder)."""
    s = s.strip().lower().replace("$", "").replace(",", "").replace("_", "")
    s = s.replace("billion", "b").replace("million", "m").replace("thousand", "k")
    m = re.match(r"^([0-9]*\.?[0-9]+)\s*([kmb]?)", s)
    if not m:
        return None
    val = float(m.group(1))
    return val * _MULT.get(m.group(2), 1.0)


def _parse_chart_series(spec):
    """'HERD expenditures | 2021=$4.1M; 2022=$4.8M; 2023=$5.3M'
    -> ('HERD expenditures', [('2021', 4.1e6), ('2022', 4.8e6), ('2023', 5.3e6)])."""
    title, _, data = spec.partition("|")
    if not data:
        title, data = "", spec
    points = []
    for chunk in data.split(";"):
        label, sep, value = chunk.partition("=")
        if not sep:
            continue
        num = _parse_money(value)
        if num is not None:
            points.append((label.strip(), num))
    return title.strip(), points


def _embed_herd_chart(doc, spec):
    """Render a `Chart:` directive as a line graph and insert it. Degrades to a
    small gray note if the data is unparseable or matplotlib is unavailable, so
    no information is lost."""
    title, points = _parse_chart_series(spec)
    if len(points) < 2:
        _small(doc.add_paragraph(), f"[HERD chart: {spec}]")
        return False
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.ticker import FuncFormatter
    except Exception:
        _small(doc.add_paragraph(), f"[HERD chart (matplotlib unavailable): {spec}]")
        return False
    labels = [p[0] for p in points]
    values = [p[1] for p in points]
    peak = max(values)
    unit, div = ("$B", 1e9) if peak >= 1e9 else ("$M", 1e6) if peak >= 1e6 else ("$K", 1e3)
    fig, ax = plt.subplots(figsize=(CHART_IN, CHART_IN * 0.55), dpi=150)
    ax.plot(labels, values, marker="o", color=MQ_BLUE, linewidth=2)
    ax.set_title(title or "HERD Research Expenditures", color=MQ_BLUE, fontsize=11)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda v, _p: f"{v / div:.1f}{unit}"))
    ax.grid(True, axis="y", linestyle=":", alpha=0.5)
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    doc.add_paragraph().add_run().add_picture(buf, width=Inches(CHART_IN))
    return True


def _find_photo(lines, start):
    """Index of the Photo: line belonging to the leader whose heading is at
    `start`, i.e. the next one before the following '### ' subhead. Flows write
    the Photo: line wherever it is convenient (usually LAST, after the bio), but
    a right-floated image anchors where it sits - so a photo written at the end
    of an entry floats beside the NEXT leader. Hoisting it to the heading makes
    placement independent of where the flow happened to put the line."""
    for j in range(start + 1, len(lines)):
        s = lines[j].strip()
        if SUBHEAD_RE.match(s):
            return None
        if PHOTO_RE.match(s):
            return j
    return None


def _place_photo(doc, val, approvals, paragraph=None):
    """Render one Photo: value: embed, or leave a small gray status note."""
    if _is_url(val):
        decision = _photo_decision(val, approvals)
        if decision == "embed":
            _embed_headshot(doc, val, paragraph=paragraph)
        elif decision == "rejected":
            _small(doc.add_paragraph(), "[headshot rejected in review]")
        else:
            _small(doc.add_paragraph(), "[headshot pending approval]")
    else:
        _small(doc.add_paragraph(), f"[headshot to add: {val}]")


def _photo_decision(url, approvals):
    if approvals is None:
        return "embed"
    status = approvals.get(url)
    if status == "approved":
        return "embed"
    if status == "rejected":
        return "rejected"
    return "pending"


def add_body(doc, body, allow_media=False, approvals=None, num_id=BULLET_NUM_GENERAL):
    lines = body.splitlines()
    consumed = set()
    pending = None   # this leader's photo, waiting for their first body paragraph
    head = None

    def flush(paragraph):
        """Anchor the waiting headshot to `paragraph`."""
        nonlocal pending
        if pending is not None and paragraph is not None:
            _place_photo(doc, pending, approvals, paragraph=paragraph)
            # keep the photo's anchor on the same page as the leader heading
            paragraph.paragraph_format.keep_with_next = True
            pending = None

    for i, raw in enumerate(lines):
        if i in consumed:
            continue
        s = raw.strip()
        if not s:
            continue
        m = SUBHEAD_RE.match(s)
        if m:
            # Previous leader had no body paragraph to hang their photo on: fall
            # back to their heading rather than let it drift onto this one.
            flush(head)
            head = doc.add_heading(m.group(1).strip(), level=2)
            # A leader's name must not strand at a page bottom with the photo
            # and experience starting overleaf.
            head.paragraph_format.keep_with_next = True
            # Hoist this leader's headshot out of wherever the flow wrote it and
            # anchor it to their FIRST body paragraph, so the heading keeps the
            # full column width and the image floats beside their own entry.
            if allow_media:
                j = _find_photo(lines, i)
                if j is not None:
                    pending = PHOTO_RE.match(lines[j].strip()).group(1)
                    consumed.add(j)
            continue
        m = PHOTO_RE.match(s)
        if m:
            # A Photo: with no leader heading above it (never hoisted); render in
            # place so the image is not silently dropped.
            if allow_media:
                _place_photo(doc, m.group(1), approvals)
            continue
        m = CHART_RE.match(s)
        if m:
            _embed_herd_chart(doc, m.group(1))
            continue
        m = profile.BULLET_RE.match(raw)
        if m:
            body_text = m.group(2).strip()
            # Provenance recedes: "Source: ..." bullets render small and gray so
            # the facts above them dominate the section.
            if body_text.startswith("Source:"):
                _small_rich(doc, body_text)
                continue
            indent = m.group(1).replace("\t", "  ")
            p = _bullet(doc, body_text, len(indent) // 2, num_id)
            flush(p)
            continue
        if s.startswith("Source:"):
            _small_rich(doc, s)
            continue
        p = doc.add_paragraph()
        _write_runs(p, s)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        flush(p)
    flush(head)   # last leader in the section


def _scrub_metadata(path, placeholder, replacement):
    """Rewrite docProps/{app,core}.xml in the saved package to drop any leftover
    template-name placeholder (invisible metadata, but tidy to remove)."""
    import zipfile
    targets = ("docProps/app.xml", "docProps/core.xml")
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in targets and placeholder.encode() in data:
                data = data.decode("utf-8").replace(placeholder, replacement).encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def load_approvals():
    if not APPROVALS_FILE.exists():
        return None
    try:
        data = json.loads(APPROVALS_FILE.read_text())
        return data if isinstance(data, dict) else None
    except (json.JSONDecodeError, OSError):
        return None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="infile",
                    default=str(ROOT / "02-assemble" / "output" / "profile-draft.md"))
    ap.add_argument("--institution", help="override institution name")
    ap.add_argument("--date", help="override date stamp (YYYY-MM-DD)")
    ap.add_argument("--no-media", action="store_true",
                    help="skip downloading/embedding leader headshots")
    args = ap.parse_args()

    src = Path(args.infile)
    if not src.exists():
        print(f"ERROR: draft not found: {src}\nRun: python 02-assemble/draft.py",
              file=sys.stderr)
        sys.exit(1)
    if not BASE_DOC.exists():
        print(f"ERROR: template base missing: {BASE_DOC}", file=sys.stderr)
        sys.exit(1)

    text = src.read_text()
    name = title_from_draft(text, args.institution)
    stamp = args.date or datetime.date.today().isoformat()
    is_short = bool(re.search(r"^#\s+Short BD Profile", text, re.MULTILINE))
    kind = "Short BD Profile" if is_short else "BD Profile"

    sections = profile.split_draft(text)
    if not sections:
        print("ERROR: no '## Section' headings found in the draft.", file=sys.stderr)
        sys.exit(1)

    approvals = None if args.no_media else load_approvals()
    doc = Document(str(BASE_DOC))
    customize_chrome(doc, name, stamp, kind)
    body, rich = reset_body(doc)

    embedded = 0
    for sect, body_text in sections:
        level = profile.heading_level(canonical(sect))
        doc.add_heading(sect, level=level)
        is_leaders = canonical(sect) == "Key Leaders"
        allow_media = is_leaders and not args.no_media
        if allow_media:
            for ln in body_text.splitlines():
                m = PHOTO_RE.match(ln.strip())
                if m and _is_url(m.group(1)) and _photo_decision(m.group(1), approvals) == "embed":
                    embedded += 1
        add_body(doc, body_text, allow_media=allow_media, approvals=approvals,
                 num_id=BULLET_NUM_LEADER if is_leaders else BULLET_NUM_GENERAL)

    if rich is not None:
        body.append(rich)

    out_dir = ROOT / "03-compile" / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{slug(kind)}_{slug(name)}_{stamp}.docx"
    doc.save(str(out))
    # docProps/app.xml (TitlesOfParts) keeps the template name as invisible
    # metadata; scrub it so the placeholder appears nowhere in the package.
    _scrub_metadata(out, PLACEHOLDER_NAME, name)

    # Count tags in the section bodies only. Scanning the whole draft also counts
    # the review-gate comment, which literally names "[verify] / [inferred]", so
    # the warning always read 2 high and disagreed with draft.py.
    tags = profile.scan_tags("\n".join(b for _, b in sections))
    print(f"Wrote {out}")
    print(f"  sections : {len(sections)}  (banner: {kind} - {name})")
    if not args.no_media:
        gate = "active" if approvals is not None else "inactive (no .headshots.json)"
        print(f"  headshots: {embedded} embedded; approval gate {gate}")
    if tags:
        print(f"  WARNING: {len(tags)} unresolved [verify]/[inferred] tags remain "
              f"in the document. Resolve in the draft and recompile if needed.")


if __name__ == "__main__":
    main()
